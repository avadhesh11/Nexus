import io
import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader
from fastapi import UploadFile

def extract_text_from_file(file: UploadFile) -> str:
    content_type = file.content_type
    filename = file.filename.lower()
    
    file_bytes = file.file.read()
    
    # Reset file pointer for potential future reads
    file.file.seek(0)
    
    # 1. Plain Text or Markdown
    if filename.endswith('.txt') or filename.endswith('.md') or content_type == "text/plain":
        return file_bytes.decode("utf-8", errors="ignore")
        
    # 2. PDF
    elif filename.endswith('.pdf') or content_type == "application/pdf":
        import pdfplumber
        text = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                
                tables = page.extract_tables()
                if tables:
                    text.append("\n--- Table Structured Data ---")
                    for table in tables:
                        for row in table:
                            row_cells = [str(cell).strip() if cell is not None else "" for cell in row]
                            if any(row_cells):
                                text.append(" | ".join(row_cells))
                    text.append("-----------------------------\n")
        return "\n".join(text)
        
    # 3. DOCX (Word Document)
    elif filename.endswith('.docx') or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(io.BytesIO(file_bytes))
        text_elements = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_elements.append(paragraph.text.strip())
                
        # Extract table text (joining cells by " | " to preserve structural association)
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate cell values if cells are merged
                row_text = []
                last_cell_text = None
                for cell in row.cells:
                    cleaned_text = cell.text.strip()
                    if cleaned_text and cleaned_text != last_cell_text:
                        row_text.append(cleaned_text)
                        last_cell_text = cleaned_text
                if row_text:
                    text_elements.append(" | ".join(row_text))
                    
        return "\n".join(text_elements)

        
    # 4. CSV or Excel
    elif filename.endswith('.csv') or content_type == "text/csv":
        df = pd.read_csv(io.BytesIO(file_bytes))
        return df.to_string(index=False)
        
    elif filename.endswith('.xlsx') or filename.endswith('.xls') or content_type in [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel"
    ]:
        df = pd.read_excel(io.BytesIO(file_bytes))
        return df.to_string(index=False)
        
    else:
        raise ValueError(f"Unsupported file format: {file.filename}")
